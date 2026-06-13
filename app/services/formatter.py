"""KDP uyumlu çıktı üretimi: DOCX / PDF / EPUB (Spec Bölüm 9).

Ağır kütüphaneler (python-docx, weasyprint, ebooklib) fonksiyon içinde lazy import edilir;
böylece bu modülü import etmek WeasyPrint sistem kütüphanelerini gerektirmez (yalnızca
generate_pdf çağrıldığında gerekir).
"""
from __future__ import annotations

import io
import os
import tempfile

import markdown as _md

# ---- KDP sabitleri (Spec Bölüm 9.1) — cm ----
KDP_FORMATS = {
    "6x9": {"width": 15.24, "height": 22.86},
    "5x8": {"width": 12.70, "height": 20.32},
    "7x10": {"width": 17.78, "height": 25.40},
    "8.5x11": {"width": 21.59, "height": 27.94},
}
MARGINS = {"top": 2.54, "bottom": 2.54, "inside": 1.91, "outside": 1.27}
BODY_FONT = "Times New Roman"
HEADING_FONT = "Arial"
BODY_SIZE = 11
HEADING_SIZES = {1: 18, 2: 14, 3: 12}
LINE_SPACING = 1.15
BLEED_MM = 3.175  # KDP 0.125 inç


def _kdp(project) -> dict:
    return KDP_FORMATS.get(project.kdp_format, KDP_FORMATS["6x9"])


def _markdown_to_html(text: str | None) -> str:
    return _md.markdown(text or "", extensions=["tables"])


def format_citation(citation, style: str = "APA") -> str:
    """Atıfı stile göre biçimler (Spec Bölüm 9.5) + doğrulanmamışsa işaretler (Bölüm 9.4)."""
    authors = citation.authors or "Anonim"
    year = citation.pub_year or "t.y."
    title = citation.raw_title
    journal = citation.journal or ""
    doi = f" https://doi.org/{citation.doi}" if citation.doi else ""
    flag = "" if citation.doi_verified else " [⚠ doğrulanamadı]"
    if style == "Chicago":
        return f'{authors}. "{title}." {journal} ({year}).{doi}{flag}'
    if style == "MLA":
        return f'{authors}. "{title}." {journal}, {year}.{doi}{flag}'
    # APA 7 (varsayılan)
    return f"{authors} ({year}). {title}. {journal}.{doi}{flag}"


# --------------------------------------------------------------------------- #
# DOCX
# --------------------------------------------------------------------------- #
def generate_docx(project, chapters, citations=None) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    doc = Document()
    section = doc.sections[0]
    fmt = _kdp(project)
    section.page_width = Cm(fmt["width"])
    section.page_height = Cm(fmt["height"])
    section.top_margin = Cm(MARGINS["top"])
    section.bottom_margin = Cm(MARGINS["bottom"])
    section.left_margin = Cm(MARGINS["inside"])
    section.right_margin = Cm(MARGINS["outside"])

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_SIZE)

    # Kapak
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(project.title)
    run.bold = True
    run.font.size = Pt(HEADING_SIZES[1])
    if project.subtitle:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.add_run(project.subtitle).italic = True

    # Telif
    doc.add_page_break()
    doc.add_paragraph(f"© {project.title}")

    # Bölümler
    for chapter in chapters:
        doc.add_page_break()
        doc.add_heading(chapter.title, level=1)
        for block in (chapter.content or "").split("\n\n"):
            block = block.strip()
            if not block:
                continue
            if block.startswith("### "):
                doc.add_heading(block[4:].strip(), level=3)
            elif block.startswith("## "):
                doc.add_heading(block[3:].strip(), level=2)
            elif block.startswith("# "):
                doc.add_heading(block[2:].strip(), level=2)
            else:
                doc.add_paragraph(block)

    # Kaynakça
    if citations:
        doc.add_page_break()
        doc.add_heading("Kaynakça", level=1)
        for citation in citations:
            doc.add_paragraph(format_citation(citation, project.citation_style))

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# --------------------------------------------------------------------------- #
# PDF (WeasyPrint)
# --------------------------------------------------------------------------- #
def generate_pdf(project, chapters, citations=None) -> bytes:
    from weasyprint import HTML

    fmt = _kdp(project)
    body = "".join(
        f"<section class='chapter'><h1>{chapter.title}</h1>{_markdown_to_html(chapter.content)}</section>"
        for chapter in chapters
    )
    refs = ""
    if citations:
        items = "".join(
            f"<li>{format_citation(c, project.citation_style)}</li>" for c in citations
        )
        refs = f"<section class='refs'><h1>Kaynakça</h1><ol>{items}</ol></section>"

    css = f"""
    @page {{
        size: {fmt['width']}cm {fmt['height']}cm;
        margin: {MARGINS['top']}cm {MARGINS['outside']}cm {MARGINS['bottom']}cm {MARGINS['inside']}cm;
        bleed: {BLEED_MM}mm;
        @bottom-center {{ content: counter(page); }}
    }}
    body {{ font-family: '{BODY_FONT}', serif; font-size: {BODY_SIZE}pt; line-height: {LINE_SPACING}; }}
    h1 {{ font-family: '{HEADING_FONT}', sans-serif; page-break-before: always; }}
    h1.title {{ page-break-before: avoid; text-align: center; }}
    """
    html = (
        "<html><head><meta charset='utf-8'><style>"
        f"{css}</style></head><body>"
        f"<h1 class='title'>{project.title}</h1>{body}{refs}</body></html>"
    )
    return HTML(string=html).write_pdf()


# --------------------------------------------------------------------------- #
# EPUB (ebooklib)
# --------------------------------------------------------------------------- #
def generate_epub(project, chapters, citations=None) -> bytes:
    from ebooklib import epub

    book = epub.EpubBook()
    book.set_identifier(str(project.id))
    book.set_title(project.title)
    book.set_language(project.language or "tr")
    book.add_author("Kitap Yazma")

    spine: list = ["nav"]
    toc: list = []
    for chapter in chapters:
        item = epub.EpubHtml(
            title=chapter.title,
            file_name=f"chap_{chapter.order_index}.xhtml",
            lang=project.language or "tr",
        )
        item.content = f"<h1>{chapter.title}</h1>{_markdown_to_html(chapter.content)}"
        book.add_item(item)
        spine.append(item)
        toc.append(item)

    if citations:
        items = "".join(f"<li>{format_citation(c, project.citation_style)}</li>" for c in citations)
        ref = epub.EpubHtml(title="Kaynakça", file_name="refs.xhtml", lang=project.language or "tr")
        ref.content = f"<h1>Kaynakça</h1><ol>{items}</ol>"
        book.add_item(ref)
        spine.append(ref)
        toc.append(ref)

    book.toc = toc
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = spine

    fd, path = tempfile.mkstemp(suffix=".epub")
    os.close(fd)
    try:
        epub.write_epub(path, book)
        with open(path, "rb") as fh:
            return fh.read()
    finally:
        os.unlink(path)
